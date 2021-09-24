import logging
import uuid
from zipfile import ZipFile
from io import BytesIO
import requests
from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfFileReader

class LearnIT:
    """A model for LearnIT, ITU's Moodle site."""
    def __init__(self, endpoint, token):
        self.endpoint = endpoint
        self.token    = token

    def create_page(self, wid, submission):
        """Create a wiki page from a PeerGrade submission."""
        params = {
            'wikiid':        wid,
            'contentformat': 'html',
            'title':         submission.title,
            'content':       submission.body
        }
        self._call_function('mod_wiki_new_page', params)

    def create_index_page(self, wid):
        """Create an index page."""
        pages = self.list_pages(wid)
        first_page = [page for page in pages if page['firstpage'] is True][0]
        content_pages = [page for page in pages if page['firstpage'] is False]
        logging.debug(f"Identified {first_page['title']} ({first_page['id']}) as first page")
        content = self.create_index_content(content_pages)
        params = {
            'pageid': first_page['id'],
            'content': content
        }
        self._call_function('mod_wiki_edit_page', params)

    def list_pages(self, wid):
        """List wiki pages."""
        logging.debug("Getting pages of wiki id %d" % wid)
        params = {
            'wikiid': wid
        }
        res = self._call_function('mod_wiki_get_subwiki_pages', params)
        return res['pages']

    def create_index_content(self, pages):
        """Just list all of them in one long index list."""
        return "<br>".join([self.wiki_linkify(page['title'])for page in pages])

    def wiki_linkify(self, page_name):
        """Turn a page name to a Wiki link."""
        return f"[[{page_name}]]"

    def _call_function(self, name, params):
        """Call the webservice endpoint, return JSON."""
        params['wsfunction']         = name
        params['wstoken']            = self.token
        params['moodlewsrestformat'] = 'json'
        return requests.post(self.endpoint, params).json()


class PeerGradeAssignment:
    """Model for a PeerGrade assignment."""
    def __init__(self, filename):
        self.filename    = filename
        self.filenames   = []
        self.submissions = []
        self.parse_submission_file()

    def parse_submission_file(self):
        """Parse a bunch of submissions to an assignment from a file."""
        logging.info("Parsing file %s" % self.filename)
        with ZipFile(self.filename) as zf:
            self.filenames = [f for f in zf.namelist() if not f.endswith('/')]
            self.read_submissions()

    def read_submissions(self):
        """Read submissions from a file."""
        for filename in self.filenames:
            with ZipFile(self.filename) as zf:
                logging.debug("Reading %s" % filename)
                with zf.open(filename) as submission_file:
                    if filename.endswith('.html'):
                        sub = SubmissionHtml(filename, submission_file.read())
                        self.submissions.append(sub)
                    elif filename.endswith('.pdf'):
                        sub = SubmissionPdf(filename, submission_file.read())
                        self.submissions.append(sub)
                    elif filename.endswith('.docx'):
                        sub = SubmissionDocx(filename, submission_file.read())
                        self.submissions.append(sub)
                    else:
                        logging.error("Unknown extension %s" % filename)
                        raise NotImplementedError


class Submission:
    """Model for a submission."""
    def __init__(self, raw, filename, body, title=None):
        self.raw = raw
        self.filename = filename
        self.body = body
        self.author, email_and_file = filename.rsplit('-', 1)
        self.email, _ = email_and_file.split('/', 1)
        self.title = title or f"📎 Oops this page needs a name {uuid.uuid1()}"
        # logging.debug("📄 %s" % raw)

    def __repr__(self):
        # return self.author
        return f"{self.__class__} {self.title} by {self.author}"


class SubmissionHtml(Submission):
    """A PeerGrade submission handed in as a HTML(ish) document."""
    def __init__(self, filename, raw):
        soup = BeautifulSoup(raw, 'html.parser')
        body = soup.body
        try:
            title = soup.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'strong', 'b']).text.strip()
        except AttributeError:
            title = None
        super().__init__(raw, filename, body, title=title)


class SubmissionDocx(Submission):
    """A PeerGrade submission handed in as a docx."""
    def __init__(self, filename, raw):
        doc = Document(BytesIO(raw))
        body = "\n".join(p.text for p in doc.paragraphs)
        title = self.infer_title()
        super().__init__(raw, filename, body, title=title)

    def infer_title(self, doc):
        """Hopefully gets the first title or heading."""
        for paragraph in doc.paragraphs:
            if paragraph.style.name in ['Title', 'Subtitle'] or paragraph.style.name.startswith('Heading'):
                return paragraph.text.strip()


class SubmissionPdf(Submission):
    """A PeerGrade submission handed in as a PDF."""
    def __init__(self, filename, raw):
        doc = PdfFileReader(BytesIO(raw))
        body = doc.getPage(0).extractText()
        super().__init__(raw, filename, body)
